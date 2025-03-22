from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, Form, File
from bson import ObjectId
from datetime import datetime
from database import projects_collection, documents_collection
from models.document import Document
from models.project import Project, ProgressEntry
from services.auth import get_current_user, get_current_admin_user
from typing import List, Optional, Dict, Any, Union
from services.cloudinary_service import cloudinary_uploader
import pandas as pd
import io
import docx
from docx.shared import Pt, Inches
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json


def get_project_or_404(project_id: str):
    project = projects_collection.find_one({"_id": ObjectId(project_id)})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


def format_date(dt):
    day = dt.day
    suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    month = dt.strftime("%B")  # Full month name
    year = dt.year
    return f"{day}{suffix} {month} {year}"


def format_number(value):
    if value is None:
        return ""
    return f"{value:,.2f}"


def parse_optional_float(value: Optional[str]) -> Optional[float]:
    if value in (None, "", "null"):
        return None
    try:
        return float(value)
    except ValueError:
        raise HTTPException(status_code=422, detail=f"Invalid float value: {value}")

def format_currency(value):
    try:
        return f"₦{float(value):,.2f}"
    except (ValueError, TypeError):
        return "N/A"


router = APIRouter()


@router.post("/", response_model=Project)
def create_project(
        project_name: str = Form(...),
        description: Optional[str] = Form(None),
        contractor: Optional[str] = Form(None),
        resident_engineer: Optional[str] = Form(None),
        progress_report: Optional[str] = Form(None),
        project_tags: Optional[str] = Form(None),
        award_date: Optional[str] = Form(None),
        contract_sum: Optional[float] = Form(None),
        duration: Optional[str] = Form(None),
        mobilisation_paid: Optional[float] = Form(None),
        interim_certificate_earned: Optional[float] = Form(None),
        remark: Optional[str] = Form(None),
        progress_of_work: Optional[str] = Form(None),
        current_user=Depends(get_current_user)
):
    try:
        # Default to None or an empty dictionary
        progress_data = None

        if progress_of_work:
            try:
                parsed_progress = json.loads(progress_of_work)  # Try parsing JSON
                if isinstance(parsed_progress, dict):
                    parsed_progress["updated_by"] = current_user.id
                    parsed_progress["updated_at"] = datetime.utcnow().isoformat()
                    progress_data = parsed_progress
                elif isinstance(parsed_progress, list):
                    progress_data = [entry if isinstance(entry, dict) else {"note": entry} for entry in parsed_progress]
            except json.JSONDecodeError:
                # If not JSON, store as a note inside a dictionary
                progress_data = {"note": progress_of_work}


        project_dict = {
            "project_name": project_name,
            "description": description,
            "contractor": contractor,
            "resident_engineer": resident_engineer,
            "progress_report": progress_report,
            "project_tags": project_tags.lower(),
            "award_date": award_date,
            "contract_sum": contract_sum,
            "duration": duration,
            "mobilisation_paid": mobilisation_paid,
            "interim_certificate_earned": interim_certificate_earned,
            "remark": remark,
            "progress_of_work": progress_data,
            "created_by": current_user.id,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        result = projects_collection.insert_one(project_dict)
        project_dict["id"] = str(result.inserted_id)
        return Project(**project_dict)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating project: {str(e)}")


# @router.post("/progress/{project_id}", response_model=dict)
# def update_project_progress(
#         project_id: str,
#         road_section: Optional[str] = Form(None),
#         building_section: Optional[str] = Form(None),
#         structural_works: Optional[str] = Form(None),
#         mse_walls: Optional[str] = Form(None),
#         drainage_works: Optional[str] = Form(None),
#         current_status: Optional[str] = Form(None),
#         dual_section: Optional[str] = Form(None),
#         single_section: Optional[str] = Form(None),
#         other_progress: Optional[str] = Form(None),
#         progress_summary: Optional[str] = Form(None),
#         current_user=Depends(get_current_user)
# ):
#     """Update project progress details separately from main project data."""
#     project = get_project_or_404(project_id)
#
#     # Create progress of work structure
#     progress_data = {
#         "road_section": road_section,
#         "building_section": building_section,
#         "structural_works": structural_works,
#         "mse_walls": mse_walls,
#         "drainage_works": drainage_works,
#         "current_status": current_status,
#         "dual_section": dual_section,
#         "single_section": single_section,
#         "other_progress": other_progress,
#         "progress_summary": progress_summary,
#         "updated_at": datetime.utcnow().isoformat()
#     }
#
#     # Remove None values
#     progress_data = {k: v for k, v in progress_data.items() if v is not None}
#
#     # Update project with progress data
#     projects_collection.update_one(
#         {"_id": ObjectId(project_id)},
#         {"$set": {"progress_of_work": progress_data, "updated_at": datetime.utcnow()}}
#     )
#
#     return {
#         "message": "Project progress updated successfully",
#         "project_id": project_id,
#         "updated_fields": list(progress_data.keys())
#     }


@router.post("/api/projects/progress/{project_id}")
async def update_project_progress(project_id: str, entry: ProgressEntry):
    # Validate project ID
    if not ObjectId.is_valid(project_id):
        raise HTTPException(status_code=400, detail="Invalid project ID")

    project_obj_id = ObjectId(project_id)

    # Check if project exists
    project = projects_collection.find_one({"_id": project_obj_id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Ensure 'progress_of_work' exists and merge updates
    existing_progress = project.get("progress_of_work", {})  # Get current progress if available
    updated_progress = {**existing_progress, **entry.progress}  # Merge new updates with existing ones

    # Update project progress_of_work
    update_result = projects_collection.update_one(
        {"_id": project_obj_id},
        {"$set": {"progress_of_work": updated_progress}}
    )

    if update_result.modified_count == 0:
        raise HTTPException(status_code=500, detail="Failed to update progress_of_work")

    return {"message": "Project progress updated successfully", "progress_of_work": updated_progress}



@router.get("/export", response_model=dict)
def export_projects():
    """Generate spreadsheet and upload to Cloudinary, including detailed progress_of_work."""
    projects = list(projects_collection.find())
    if not projects:
        raise HTTPException(status_code=404, detail="No projects found")

    project_data = []

    for idx, proj in enumerate(projects):
        # Extract progress details
        progress_details = ""
        if isinstance(proj.get("progress_of_work"), dict):
            sections = []
            for section, value in proj["progress_of_work"].items():
                if value:
                    sections.append(f"{section.replace('_', ' ').title()}: {value}")
            progress_details = "\n".join(sections) if sections else "No progress reported"
        else:
            progress_details = str(proj.get("progress_of_work", "No progress reported"))

        project_entry = {
            "S/N": idx + 1,
            "Project Name": proj.get("project_name", "N/A"),
            "Contractor": proj.get("contractor", "N/A"),
            "Resident Engineer": proj.get("resident_engineer", "N/A"),
            "Progress Report": proj.get("progress_report", "N/A"),
            "Project Tags": proj.get("project_tags", "N/A"),
            "Award Date": proj.get("award_date", "N/A"),
            "Contract Sum": format_currency(proj.get("contract_sum", "N/A")),
            "Duration": proj.get("duration", "N/A"),
            "Mobilisation Paid": format_currency(proj.get("mobilisation_paid", "N/A")),
            "Interim Certificate Earned": format_currency(proj.get("interim_certificate_earned", "N/A")),
            "Progress of Work": progress_details,
            "Remark": proj.get("remark", "N/A"),
        }
        project_data.append(project_entry)

    # Export to Excel
    df = pd.DataFrame(project_data)
    output_excel = io.BytesIO()
    df.to_excel(output_excel, index=False, engine="openpyxl")
    output_excel.seek(0)

    # Upload Excel file to Cloudinary
    output_excel_file = UploadFile(filename="projects_export.xlsx", file=output_excel)
    upload_result_excel = cloudinary_uploader.upload(output_excel_file.file, folder="project_exports")

    # Create Word document with formatted progress details
    doc = docx.Document()
    now = datetime.now()
    formatted_date_str = format_date(now)
    doc.add_heading(f"PROJECT PROGRESS REPORT AS OF {formatted_date_str}", level=1)

    table = doc.add_table(rows=1, cols=len(project_data[0]))
    table.style = "Table Grid"

    for row in table.rows:
        for cell in row.cells:
            cell.width = docx.shared.Inches(1.5)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = docx.shared.Pt(6)


    hdr_cells = table.rows[0].cells
    for i, key in enumerate(project_data[0].keys()):
        hdr_cells[i].text = key
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Add rows
    for proj in project_data:
        row_cells = table.add_row().cells
        for i, (key, value) in enumerate(proj.items()):
            if key == "Progress of Work":
                paragraph = row_cells[i].paragraphs[0]
                lines = str(value).split("\n")
                if lines:
                    paragraph.text = lines[0]
                    for line in lines[1:]:
                        if line.strip():
                            paragraph.add_run("\n" + line)
            else:
                row_cells[i].text = str(value)

    # Save and upload Word document
    output_word = io.BytesIO()
    doc.save(output_word)
    output_word.seek(0)

    output_word_file = UploadFile(filename="projects_export.docx", file=output_word)
    upload_result_word = cloudinary_uploader.upload(output_word_file.file, folder="project_exports")

    return {
        "spreadsheet_url": upload_result_excel,
        "word_doc_url": upload_result_word,
        "row_count": len(project_data),
        "column_count": len(project_data[0])
    }


@router.get("/export/ongoing", response_model=dict)
def export_ongoing_projects():
    """Generate spreadsheet for ongoing projects and upload to Cloudinary."""
    projects = list(projects_collection.find({"project_tags": "ongoing"}))
    if not projects:
        raise HTTPException(status_code=404, detail="No ongoing projects found")

    project_data = []

    for idx, proj in enumerate(projects):
        progress_details = ""
        if isinstance(proj.get("progress_of_work"), dict):
            sections = [f"{key.replace('_', ' ').title()}: {value}" for key, value in proj["progress_of_work"].items()
                        if value]
            progress_details = "\n".join(sections) if sections else "No progress reported"
        else:
            progress_details = str(proj.get("progress_of_work", "No progress reported"))

        project_entry = {
            "S/N": idx + 1,
            "Project Name": proj.get("project_name", "N/A"),
            "Contractor": proj.get("contractor", "N/A"),
            "Resident Engineer": proj.get("resident_engineer", "N/A"),
            "Progress Report": proj.get("progress_report", "N/A"),
            # "Progress of Work": progress_details
        }
        project_data.append(project_entry)

    # Export to Excel
    df = pd.DataFrame(project_data)
    output_excel = io.BytesIO()
    df.to_excel(output_excel, index=False, engine="openpyxl")
    output_excel.seek(0)

    # Create Word document
    doc = docx.Document()
    now = datetime.now()
    formatted_date_str = format_date(now)
    doc.add_heading(f"ONGOING PROJECTS PROGRESS REPORT AS OF {formatted_date_str}", level=1)

    table = doc.add_table(rows=1, cols=len(project_data[0]))
    table.style = "Table Grid"

    for row in table.rows:
        for cell in row.cells:
            cell.width = docx.shared.Inches(1.5)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = docx.shared.Pt(6)



    hdr_cells = table.rows[0].cells
    for i, key in enumerate(project_data[0].keys()):
        hdr_cells[i].text = key
        hdr_cells[i].paragraphs[0].runs[0].bold = True


    for proj in project_data:
        row_cells = table.add_row().cells
        for i, (key, value) in enumerate(proj.items()):
            if key == "Progress of Work":
                paragraph = row_cells[i].paragraphs[0]
                lines = str(value).split("\n")
                if lines:
                    paragraph.text = lines[0]
                    for line in lines[1:]:
                        if line.strip():
                            paragraph.add_run("\n" + line)
            else:
                row_cells[i].text = str(value)

    # Save and upload Word document
    output_word = io.BytesIO()
    doc.save(output_word)
    output_word.seek(0)

    # Upload to Cloudinary
    upload_result_excel = cloudinary_uploader.upload(output_excel, folder="project_exports")
    upload_result_word = cloudinary_uploader.upload(output_word, folder="project_exports")

    return {
        "spreadsheet_url": upload_result_excel,
        "word_document_url": upload_result_word
    }


@router.get("/", response_model=List[Project])
def get_projects():
    """Retrieve all projects sorted by creation date."""
    projects = projects_collection.find().sort("created_at", -1)

    def sanitize_data(data):
        """Recursively replace NaN values with None."""
        import math
        if isinstance(data, float) and math.isnan(data):
            return None
        elif isinstance(data, dict):
            return {k: sanitize_data(v) for k, v in data.items()}
        elif isinstance(data, list):
            return [sanitize_data(v) for v in data]
        return data

    return [
        Project(
            **sanitize_data(
                {
                    **project,
                    "id": str(project["_id"]),
                    "progress_of_work": project.get("progress_of_work", {}) if isinstance(
                        project.get("progress_of_work"), dict) else {}
                }
            )
        )
        for project in projects
    ]


@router.get("/id/{project_id}", response_model=Project)
def get_project_by_id(project_id: str):
    """Retrieve a project by its ID."""
    project = projects_collection.find_one({"_id": ObjectId(project_id)})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return Project(**project, id=str(project["_id"]))


@router.get("/progress/{project_id}", response_model=Dict[str, Any])
def get_project_progress(project_id: str):
    """Get detailed progress information for a project."""
    project = get_project_or_404(project_id)

    progress = project.get("progress_of_work", {})
    if not progress:
        progress = {}

    return {
        "project_id": project_id,
        "project_name": project.get("project_name", ""),
        "progress_data": progress
    }


@router.get("/name/{project_name}", response_model=List[Project])
def get_project_by_name(project_name: str):
    """Retrieve projects by name using regex matching."""
    projects = projects_collection.find(
        {"project_name": {"$regex": project_name, "$options": "i"}}
    )
    projects_list = list(projects)
    if not projects_list:
        raise HTTPException(status_code=404, detail="No projects found with this name")
    return [
        Project(**{**project, "id": str(project["_id"])} if "_id" in project else project)
        for project in projects_list
    ]


@router.get("/recent", response_model=List[Project])
def get_recent_projects(limit: int = 5, user=Depends(get_current_user)):
    """Retrieve the most recently uploaded projects."""
    projects = list(projects_collection.find().sort([("created_at", -1)]).limit(limit))
    return [Project(**{**proj, "id": str(proj.pop("_id"))}) for proj in projects]


@router.get("/{project_id}/documents", response_model=List[Document])
def get_project_documents(project_id: str):
    """Retrieve all documents associated with a specific project."""
    documents = list(documents_collection.find({"project_id": project_id}))
    if not documents:
        raise HTTPException(status_code=404, detail="No documents found for this project")
    return [Document(**{**doc, "id": str(doc["_id"]), "_id": str(doc["_id"])}) for doc in documents]


@router.put("/projects/{project_id}")
def update_project(
    project_id: str,
    project_name: Optional[str] = Form(None),
    contractor: Optional[str] = Form(None),
    resident_engineer: Optional[str] = Form(None),
    progress_report: Optional[str] = Form(None),
    project_tags: Optional[str] = Form(None),
    award_date: Optional[str] = Form(None),
    contract_sum: Optional[str] = Form(None),
    duration: Optional[str] = Form(None),
    mobilisation_paid: Optional[str] = Form(None),
    interim_certificate_earned: Optional[str] = Form(None),
    remark: Optional[str] = Form(None),
    current_user: dict = Depends(get_current_user),
):
    project = get_project_or_404(project_id)

    update_data = {}

    # Only add fields to update_data if provided
    if project_name:
        update_data["project_name"] = project_name
    if contractor:
        update_data["contractor"] = contractor
    if resident_engineer:
        update_data["resident_engineer"] = resident_engineer
    if progress_report:
        update_data["progress_report"] = progress_report
    if project_tags:
        update_data["project_tags"] = project_tags.lower()
    if award_date:
        update_data["award_date"] = award_date
        # try:
        #     update_data["award_date"] = datetime.strptime(award_date, "%Y-%m-%d")
        # except ValueError:
        #     raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")
    if duration:
        update_data["duration"] = duration
    if remark:
        update_data["remark"] = remark

    # Convert float fields only if they have valid values
    if contract_sum:
        update_data["contract_sum"] = parse_optional_float(contract_sum)
    if mobilisation_paid:
        update_data["mobilisation_paid"] = parse_optional_float(mobilisation_paid)
    if interim_certificate_earned:
        update_data["interim_certificate_earned"] = parse_optional_float(interim_certificate_earned)

    if not update_data:
        raise HTTPException(status_code=400, detail="No valid fields provided for update")

    update_data["updated_at"] = datetime.utcnow()

    # Update only the provided fields
    projects_collection.update_one({"_id": ObjectId(project_id)}, {"$set": update_data})

    return {"message": "Project updated successfully", "updated_fields": list(update_data.keys())}

@router.delete("/{project_id}", response_model=dict)
def delete_project(project_id: str, current_user=Depends(get_current_admin_user)):
    """Delete a project (Admin only)."""
    project = projects_collection.find_one({"_id": ObjectId(project_id)})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if "file_url" in project:
        try:
            cloudinary_uploader.delete(project["file_url"].split("/")[-1].split(".")[0])
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error deleting file: {str(e)}")

    projects_collection.delete_one({"_id": ObjectId(project_id)})
    return {"message": "Project deleted successfully"}